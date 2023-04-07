import csv

from docx import Document
from docx.table import Table
from docx.shared import Inches


# read csv for rule set data
def readRuleSetCSV():
    with open('/Users/rishabhjain/Downloads/rules.csv', 'r') as f:
        reader = csv.reader(f)
        rules = list(reader)
    return rules

# define a class for rule set
class RuleSet:
    clarify = ''
    fieldName = ''
    fieldDescription = ''
    alternativeName = ''
    fieldType = ''
    cerLocation = ''
    locationType = ''
    sscpDestinationPartA = ''
    sscpDestinationPartB = ''
    languageConversionNeededForPartB = ''
    note = ''
    comment = ''

    def __init__(self, clarify, fieldName, fieldDescription, alternativeName, fieldType, cerLocation, locationType,
                 sscpDestinationPartA, sscpDestinationPartB, languageConversionNeededForPartB, note, comment):
        self.clarify = clarify
        self.fieldName = fieldName
        self.fieldDescription = fieldDescription
        self.alternativeName = alternativeName
        self.fieldType = fieldType
        self.cerLocation = cerLocation
        self.locationType = locationType
        self.sscpDestinationPartA = sscpDestinationPartA
        self.sscpDestinationPartB = sscpDestinationPartB
        self.languageConversionNeededForPartB = languageConversionNeededForPartB
        self.note = note
        self.comment = comment


# parse csv and create rule set object
def parseRuleSetCSV():
    rules = readRuleSetCSV()
    ruleSetList = []

    # ignore first row
    rules.pop(0)

    for rule in rules:
        ruleSet = RuleSet(rule[0], rule[1], rule[2], rule[3], rule[4], rule[5], rule[6], rule[7], rule[8], rule[9],
                          rule[10], rule[11])
        ruleSetList.append(ruleSet)
    return ruleSetList

# print all rule set
def printRuleSet():
    ruleSetList = parseRuleSetCSV()
    for ruleSet in ruleSetList:
        print(ruleSet.clarify)
        print(ruleSet.fieldName)
        print(ruleSet.fieldDescription)
        print(ruleSet.alternativeName)
        print(ruleSet.fieldType)
        print(ruleSet.cerLocation)
        print(ruleSet.locationType)
        print(ruleSet.sscpDestinationPartA)
        print(ruleSet.sscpDestinationPartB)
        print(ruleSet.languageConversionNeededForPartB)
        print(ruleSet.note)
        print(ruleSet.comment)

#printRuleSet()

# read input docx file
def readInputDocxFile():
    document = Document('/Users/rishabhjain/Downloads/OneDrive_1_4-6-2023/Word-word-conversion/Artifacts/input.docx')
    return document

# print all heading 1 and heading 2
def printHeading(document):
    for para in document.paragraphs:
        if para.style.name == 'Heading 1':
            print(para.text)
        if para.style.name == 'Heading 2':
            print(para.text)

#printHeading(readInputDocxFile())

# read template docx file
def readTemplateDocxFile():
    document = Document('/Users/rishabhjain/Downloads/OneDrive_1_4-6-2023/Word-word-conversion/Artifacts/SSCP_Template_Rev10.docx')
    return document

#printHeading(readTemplateDocxFile())

# save the updated template docx file
def saveDocxFile(document):
    document.save('/Users/rishabhjain/Downloads/OneDrive_1_4-6-2023/Word-word-conversion/Artifacts/output.docx')


# for each rule set locate data in input docx file, update template docx file
def enforceRules():
    ruleSetList = parseRuleSetCSV()
    inputDocument = readInputDocxFile()
    templateDocument = readTemplateDocxFile()

    for ruleSet in ruleSetList:
        # if location type starts from "table(" then it is a table
        if ruleSet.locationType.startswith("Table ("):
            # get table name from location type
            # remove "Table (" and ")"
            tableName = ruleSet.locationType[7:-1].split(")")[0]
            # find table number by locating the table name
            tableNumber = 0
            for para in inputDocument.paragraphs:
                # get caption of table
                # get table by caption
                if tableName in para.text:
                    tableNumber = para._element
                    
                    tableNumber = tableNumber.getparent()
                    print(tableNumber)
            
            # get table from input docx file
            table = inputDocument.tables[tableNumber]

            # print table heading column 0
            for row in table.rows:
                print(row)


    #saveDocxFile(templateDocument)


enforceRules()